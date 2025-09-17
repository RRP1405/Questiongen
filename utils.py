import pdfplumber
from docx import Document as DocxDocument
from docx.shared import Pt
import docx
import os
import random
import re

# -----------------------
# Helpers: read files
# -----------------------
def extract_text(file_path):
    file_path = file_path.replace("\\", "/")
    if file_path.lower().endswith(".pdf"):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for p in pdf.pages:
                page_text = p.extract_text() or ""
                text += page_text + "\n"
        return text
    elif file_path.lower().endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif file_path.lower().endswith(".docx"):
        doc = docx.Document(file_path)
        full = []
        for para in doc.paragraphs:
            full.append(para.text)
        return "\n".join(full)
    else:
        # unknown type: try reading as text
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except:
            return ""

# -----------------------
# Extract topic lines
# -----------------------
def extract_topic_lines(text):
    # Break into lines, clean, remove very short noise lines
    lines = []
    for raw in text.splitlines():
        l = raw.strip()
        # Remove common bullet characters from start
        l = re.sub(r'^[\-\*\•\d\.\)\s]+', '', l).strip()
        if len(l) >= 3:
            lines.append(l)
    # Deduplicate while preserving order
    seen = set()
    topics = []
    for l in lines:
        if l.lower() not in seen:
            seen.add(l.lower())
            topics.append(l)
    return topics

# -----------------------
# Make question templates from topics
# -----------------------
def make_question_templates(topics):
    # If syllabus is long, topics list will be long
    # We'll create different pools by difficulty mark (1,2,5,10)
    one_mark_choose = []
    one_mark_fill = []
    two_mark = []
    five_mark = []
    ten_mark = []

    for t in topics:
        # sanitize topic short version
        short = t
        # 1 mark choose: simple MCQ template (options are generic)
        mcq_q = f"Which of the following is correct about \"{short}\"?"
        # create 4 generic options (user can edit later)
        options = [
            f"{short}",
            f"Not related to {short}",
            f"Partially related to {short}",
            "None of the above"
        ]
        one_mark_choose.append( (mcq_q, options) )

        # 1 mark fill: statement with a missing keyword (first noun-like word)
        words = re.findall(r'\w+', short)
        if len(words) >= 1:
            blank_word = words[0]
            fill_q = short.replace(blank_word, "___")
        else:
            fill_q = f"Fill in the blank: ___ about {short}"
        one_mark_fill.append(fill_q)

        # 2 mark: definition/short explanation
        two_mark.append(f"Define / Explain briefly: {short}")

        # 5 mark: compare/explain in brief or either-or style prompt
        five_mark.append(f"Explain in detail (5m): {short}")

        # 10 mark: long answer / discuss
        ten_mark.append(f"Discuss in detail (10m): {short}")

    # If any pool is empty, add fallback templates
    if not one_mark_choose:
        one_mark_choose = [("Which one is true?", ["A", "B", "C", "D"])]
    if not one_mark_fill:
        one_mark_fill = ["Fill in the blank: ___"]
    if not two_mark:
        two_mark = ["Explain: Topic"]
    if not five_mark:
        five_mark = ["Explain in detail: Topic"]
    if not ten_mark:
        ten_mark = ["Discuss: Topic"]

    return {
        "choose": one_mark_choose,
        "fill": one_mark_fill,
        "two": two_mark,
        "five": five_mark,
        "ten": ten_mark
    }

# -----------------------
# Selection based on pattern
# -----------------------
def select_questions(pools, paper_type):
    """
    paper_type is "50" or "75"
    returns dict with sections
    """
    random.seed()  # system random

    selected = {
        "one_choose": [],
        "one_fill": [],
        "two": [],
        "five_pairs": [],  # list of tuples (qA, qB) for either/or
        "ten": []
    }

    if paper_type == "50":
        # 6 one-mark (3 choose, 3 fill)
        selected["one_choose"] = random.sample(pools["choose"], min(3, len(pools["choose"])))
        selected["one_fill"] = random.sample(pools["fill"], min(3, len(pools["fill"])))
        # 2 two-mark
        selected["two"] = random.sample(pools["two"], min(2, len(pools["two"])))
        # 8 five-mark in either-or => create 4 pairs (each pair has two alternatives)
        five_needed = 8
        # if pools["five"] big, sample 8, split into 4 pairs
        many_five = random.sample(pools["five"], min(five_needed, len(pools["five"])))
        # if less than needed, allow repetition
        while len(many_five) < five_needed:
            many_five += random.sample(pools["five"], min(five_needed - len(many_five), len(pools["five"])))
        # pair them
        for i in range(0, five_needed, 2):
            selected["five_pairs"].append( (many_five[i], many_five[i+1]) )
        # 3 ten-mark
        selected["ten"] = random.sample(pools["ten"], min(3, len(pools["ten"])))
    else:  # 75
        # 10 one-mark (5 choose, 5 fill)
        selected["one_choose"] = random.sample(pools["choose"], min(5, len(pools["choose"])))
        selected["one_fill"] = random.sample(pools["fill"], min(5, len(pools["fill"])))
        # 5 two-mark
        selected["two"] = random.sample(pools["two"], min(5, len(pools["two"])))
        # 10 five-mark -> 5 pairs
        five_needed = 10
        many_five = random.sample(pools["five"], min(five_needed, len(pools["five"])))
        while len(many_five) < five_needed:
            many_five += random.sample(pools["five"], min(five_needed - len(many_five), len(pools["five"])))
        for i in range(0, five_needed, 2):
            selected["five_pairs"].append( (many_five[i], many_five[i+1]) )
        # 5 ten-mark
        selected["ten"] = random.sample(pools["ten"], min(5, len(pools["ten"])))

    return selected

# -----------------------
# Create docx formatted paper
# -----------------------
def create_doc(selected_qs, paper_type, output_folder):
    doc = DocxDocument()
    # Basic style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.add_heading(f"Question Paper - {paper_type} Marks", level=1)
    doc.add_paragraph("Time: ____    Max Marks: " + str(paper_type))
    doc.add_paragraph("Name: __________  Reg No: ____")
    doc.add_paragraph("")

    # Section A - One mark
    doc.add_heading("Section A: One mark questions", level=2)
    # choose
    if selected_qs["one_choose"]:
        doc.add_paragraph("A1. Choose the correct option: (1 × {})".format(len(selected_qs["one_choose"])))
        idx = 1
        for q, opts in selected_qs["one_choose"]:
            p = doc.add_paragraph(f"{idx}. {q}")
            # options
            letters = ['A', 'B', 'C', 'D']
            for i, opt in enumerate(opts):
                doc.add_paragraph(f"    {letters[i % 4]}. {opt}")
            idx += 1
        doc.add_paragraph("")

    # fill
    if selected_qs["one_fill"]:
        doc.add_paragraph("A2. Fill in the blanks: (1 × {})".format(len(selected_qs["one_fill"])))
        idx = 1
        for q in selected_qs["one_fill"]:
            doc.add_paragraph(f"{idx}. {q}")
            idx += 1
        doc.add_paragraph("")

    # Section B - Two mark
    doc.add_heading("Section B: Two mark questions", level=2)
    idx = 1
    for q in selected_qs["two"]:
        doc.add_paragraph(f"{idx}. {q}   (2)")
        idx += 1
    doc.add_paragraph("")

    # Section C - Five mark (either-or pairs)
    doc.add_heading("Section C: Five mark questions (Answer either-or)", level=2)
    idx = 1
    for pair in selected_qs["five_pairs"]:
        q1, q2 = pair
        doc.add_paragraph(f"{idx}. a) {q1}   (5)\n    OR\n    b) {q2}   (5)")
        idx += 1
    doc.add_paragraph("")

    # Section D - Ten mark
    doc.add_heading("Section D: Ten mark questions", level=2)
    idx = 1
    for q in selected_qs["ten"]:
        doc.add_paragraph(f"{idx}. {q}   (10)")
        idx += 1

    # Save file
    os.makedirs(output_folder, exist_ok=True)
    fname = f"question_paper_{paper_type}_{random.randint(1000,9999)}.docx"
    outpath = os.path.join(output_folder, fname)
    doc.save(outpath)
    return outpath

# -----------------------
# Main pipeline
# -----------------------
def process_syllabus(file_path, paper_type, output_folder):
    text = extract_text(file_path)
    topics = extract_topic_lines(text)
    pools = make_question_templates(topics)
    selected = select_questions(pools, str(paper_type))
    out = create_doc(selected, paper_type, output_folder)
    return out